import PyPDF2
from pathlib import Path

try:
    from docx import Document as WordDocument
except ImportError:  # pragma: no cover - handled via setup
    WordDocument = None


class DocumentHandler:
    """Handles PDF, Word, and text file processing"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page.extract_text()
        
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return text
    
    @staticmethod
    def extract_text_from_txt(txt_path):
        """Extract text from text file"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # Try different encoding
            with open(txt_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
        
        return text

    @staticmethod
    def extract_text_from_docx(docx_path):
        """Extract text from Word document"""
        if WordDocument is None:
            raise ImportError("python-docx is required for Word document processing. Install with `pip install python-docx`.")

        try:
            doc = WordDocument(docx_path)
            parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Include table text if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        parts.append(" | ".join(row_text))

            text = "\n".join(parts).strip()
            if not text:
                text = " ".join(run.text for para in doc.paragraphs for run in para.runs)

            return text or "No readable text found in the Word document."

        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")
    
    @staticmethod
    def extract_text(file_path):
        """Extract text based on file type"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return DocumentHandler.extract_text_from_pdf(file_path)
        elif suffix == '.txt':
            return DocumentHandler.extract_text_from_txt(file_path)
        elif suffix in ('.docx', '.doc'):
            return DocumentHandler.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
