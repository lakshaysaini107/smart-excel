# PDF and Document Processor for Smart Excel Assistant

import io
import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from datetime import datetime
import textstat

# PDF Processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word Document Processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Text Analysis
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

logger = logging.getLogger(__name__)

class PDFDocumentProcessor:
    """Advanced PDF and Document Processor with text extraction and analysis"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': 'PDF Document',
            '.txt': 'Text Document', 
            '.docx': 'Word Document'
        }
        
        # Initialize NLTK if available
        if NLTK_AVAILABLE:
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
            except Exception as e:
                logger.warning(f"NLTK setup warning: {e}")
    
    def process_document(self, uploaded_file) -> Dict[str, Any]:
        """
        Main method to process uploaded documents
        """
        try:
            file_extension = Path(uploaded_file.name).suffix.lower()
            
            if file_extension not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_extension}. Supported: {", ".join(self.supported_formats.keys())}'
                }
            
            logger.info(f"Processing document: {uploaded_file.name} ({file_extension})")
            
            # Extract text based on file type
            if file_extension == '.pdf':
                extracted_text = self._extract_pdf_text(uploaded_file)
            elif file_extension == '.docx':
                extracted_text = self._extract_docx_text(uploaded_file)
            elif file_extension == '.txt':
                extracted_text = self._extract_txt_text(uploaded_file)
            else:
                return {
                    'success': False,
                    'error': f'Processing for {file_extension} not implemented'
                }
            
            if not extracted_text or not extracted_text.strip():
                return {
                    'success': False,
                    'error': 'No text could be extracted from the document'
                }
            
            # Clean and process the extracted text
            cleaned_text = self._clean_text(extracted_text)
            
            # Generate document metrics
            metrics = self._analyze_document_metrics(cleaned_text)
            
            # Structure the text for analysis
            structured_content = self._structure_content(cleaned_text)
            
            return {
                'success': True,
                'filename': uploaded_file.name,
                'file_type': self.supported_formats[file_extension],
                'file_size': len(uploaded_file.getvalue()),
                'raw_text': extracted_text,
                'cleaned_text': cleaned_text,
                'structured_content': structured_content,
                'metrics': metrics,
                'extraction_method': file_extension,
                'processed_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                'success': False,
                'error': f'Document processing failed: {str(e)}',
                'filename': uploaded_file.name if uploaded_file else 'Unknown'
            }
    
    def _extract_pdf_text(self, uploaded_file) -> str:
        """Extract text from PDF using multiple methods"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available. Install PyPDF2 and pdfplumber.")
        
        text_content = ""
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as page_error:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                        continue
            
            # If pdfplumber didn't work well, try PyPDF2
            if not text_content.strip():
                uploaded_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as page_error:
                        logger.warning(f"PyPDF2 failed to extract text from page {page_num + 1}: {page_error}")
                        continue
                        
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise Exception(f"Could not extract text from PDF: {str(e)}")
        
        return text_content
    
    def _extract_docx_text(self, uploaded_file) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("Word document processing not available. Install python-docx.")
        
        try:
            uploaded_file.seek(0)
            doc = DocxDocument(uploaded_file)
            
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract tables if present
            if doc.tables:
                text_content += "\n--- Tables ---\n"
                for table_num, table in enumerate(doc.tables):
                    text_content += f"\nTable {table_num + 1}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_content += row_text + "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise Exception(f"Could not extract text from Word document: {str(e)}")
    
    def _extract_txt_text(self, uploaded_file) -> str:
        """Extract text from plain text file"""
        try:
            uploaded_file.seek(0)
            content = uploaded_file.read()
            
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
                try:
                    if isinstance(content, bytes):
                        return content.decode(encoding)
                    else:
                        return str(content)
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            # If all encodings fail, use error handling
            if isinstance(content, bytes):
                return content.decode('utf-8', errors='replace')
            else:
                return str(content)
                
        except Exception as e:
            logger.error(f"TXT text extraction failed: {e}")
            raise Exception(f"Could not extract text from text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        try:
            # Remove excessive whitespace
            text = re.sub(r'\n\s*\n', '\n\n', text)
            text = re.sub(r'[ \t]+', ' ', text)
            
            # Remove page markers
            text = re.sub(r'--- Page \d+ ---', '', text)
            
            # Fix common OCR issues
            text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between words
            text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Space after punctuation
            
            # Remove excessive punctuation
            text = re.sub(r'[.]{3,}', '...', text)
            text = re.sub(r'[-]{3,}', '---', text)
            
            # Clean up spacing
            text = text.strip()
            
            return text
            
        except Exception as e:
            logger.warning(f"Text cleaning failed: {e}")
            return text  # Return original text if cleaning fails
    
    def _analyze_document_metrics(self, text: str) -> Dict[str, Any]:
        """Analyze document metrics and readability"""
        try:
            metrics = {}
            
            # Basic counts
            metrics['character_count'] = len(text)
            metrics['character_count_no_spaces'] = len(text.replace(' ', ''))
            
            # Word count
            words = text.split()
            metrics['word_count'] = len(words)
            
            # Sentence count
            if NLTK_AVAILABLE:
                try:
                    sentences = sent_tokenize(text)
                    metrics['sentence_count'] = len(sentences)
                except:
                    # Fallback method
                    sentences = re.split(r'[.!?]+', text)
                    metrics['sentence_count'] = len([s for s in sentences if s.strip()])
            else:
                sentences = re.split(r'[.!?]+', text)
                metrics['sentence_count'] = len([s for s in sentences if s.strip()])
            
            # Paragraph count
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            metrics['paragraph_count'] = len(paragraphs)
            
            # Average metrics
            if metrics['sentence_count'] > 0:
                metrics['avg_words_per_sentence'] = round(metrics['word_count'] / metrics['sentence_count'], 2)
            else:
                metrics['avg_words_per_sentence'] = 0
            
            if metrics['paragraph_count'] > 0:
                metrics['avg_sentences_per_paragraph'] = round(metrics['sentence_count'] / metrics['paragraph_count'], 2)
            else:
                metrics['avg_sentences_per_paragraph'] = 0
            
            # Reading time (average 200 words per minute)
            metrics['reading_time_minutes'] = round(metrics['word_count'] / 200, 1)
            
            # Readability scores using textstat
            try:
                # Flesch Reading Ease (0-100, higher = easier)
                metrics['flesch_reading_ease'] = round(textstat.flesch_reading_ease(text), 2)
                
                # Flesch-Kincaid Grade Level
                metrics['flesch_kincaid_grade'] = round(textstat.flesch_kincaid_grade(text), 2)
                
                # Automated Readability Index
                metrics['automated_readability_index'] = round(textstat.automated_readability_index(text), 2)
                
                # Overall complexity score (1=easy, 2=medium, 3=hard)
                if metrics['flesch_reading_ease'] > 60:
                    metrics['complexity_score'] = 1  # Easy
                elif metrics['flesch_reading_ease'] > 30:
                    metrics['complexity_score'] = 2  # Medium
                else:
                    metrics['complexity_score'] = 3  # Hard
                    
            except Exception as readability_error:
                logger.warning(f"Readability analysis failed: {readability_error}")
                metrics['complexity_score'] = 2  # Default to medium
            
            # Language detection (simple heuristic)
            common_english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were']
            text_lower = text.lower()
            english_word_count = sum(1 for word in common_english_words if word in text_lower)
            metrics['likely_english'] = english_word_count >= 5
            
            return metrics
            
        except Exception as e:
            logger.error(f"Document metrics analysis failed: {e}")
            return {
                'character_count': len(text),
                'word_count': len(text.split()),
                'complexity_score': 2,
                'error': str(e)
            }
    
    def _structure_content(self, text: str) -> Dict[str, Any]:
        """Structure the content into sections and key elements"""
        try:
            structured = {
                'sections': [],
                'key_phrases': [],
                'entities': {
                    'numbers': [],
                    'dates': [],
                    'urls': [],
                    'emails': []
                }
            }
            
            # Split into sections (basic approach)
            # Look for common section headers
            section_patterns = [
                r'\n\s*(?:Chapter|Section|Part)\s+\d+[:\s]',
                r'\n\s*\d+\.\s+[A-Z][a-zA-Z\s]+\n',
                r'\n\s*[A-Z][A-Z\s]{3,}\n',  # ALL CAPS headers
                r'\n\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*\n'  # Title Case headers
            ]
            
            current_section = {'title': 'Introduction', 'content': ''}
            sections = [current_section]
            
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                # Check if this looks like a section header
                is_header = False
                for pattern in section_patterns:
                    if re.match(pattern, '\n' + paragraph):
                        is_header = True
                        break
                
                if is_header and len(paragraph) < 200:  # Headers are usually short
                    # Start new section
                    current_section = {'title': paragraph, 'content': ''}
                    sections.append(current_section)
                else:
                    current_section['content'] += paragraph + '\n\n'
            
            structured['sections'] = sections
            
            # Extract entities
            # Numbers (including percentages, currency)
            number_patterns = [
                r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b',  # Regular numbers with commas
                r'\b\d+(?:\.\d+)?%\b',  # Percentages
                r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?\b',  # Currency
                r'\b\d+(?:\.\d+)?(?:\s*(?:million|billion|thousand|k|M|B))?\b'  # Large numbers
            ]
            
            for pattern in number_patterns:
                matches = re.findall(pattern, text)
                structured['entities']['numbers'].extend(matches)
            
            # Dates
            date_patterns = [
                r'\b\d{1,2}/\d{1,2}/\d{4}\b',  # MM/DD/YYYY
                r'\b\d{1,2}-\d{1,2}-\d{4}\b',  # MM-DD-YYYY
                r'\b\d{4}-\d{2}-\d{2}\b',      # YYYY-MM-DD
                r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b',  # Month DD, YYYY
                r'\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b'   # DD Month YYYY
            ]
            
            for pattern in date_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                structured['entities']['dates'].extend(matches)
            
            # URLs
            url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+|www\.[^\s<>"{}|\\^`\[\]]+'
            urls = re.findall(url_pattern, text, re.IGNORECASE)
            structured['entities']['urls'] = urls
            
            # Email addresses
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, text)
            structured['entities']['emails'] = emails
            
            # Extract key phrases (simple approach using capitalized phrases)
            key_phrase_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b'
            potential_phrases = re.findall(key_phrase_pattern, text)
            
            # Filter out common words and keep significant phrases
            common_words = {'The', 'This', 'That', 'These', 'Those', 'And', 'Or', 'But', 'In', 'On', 'At', 'To', 'For', 'Of', 'With', 'By'}
            key_phrases = [phrase for phrase in potential_phrases 
                          if phrase not in common_words and len(phrase) > 3]
            
            # Get unique phrases and limit to most frequent ones
            from collections import Counter
            phrase_counts = Counter(key_phrases)
            structured['key_phrases'] = [phrase for phrase, count in phrase_counts.most_common(20)]
            
            return structured
            
        except Exception as e:
            logger.error(f"Content structuring failed: {e}")
            return {
                'sections': [{'title': 'Full Document', 'content': text}],
                'key_phrases': [],
                'entities': {'numbers': [], 'dates': [], 'urls': [], 'emails': []},
                'error': str(e)
            }
    
    def prepare_for_summarization(self, processed_content: Dict) -> List[str]:
        """Prepare processed content for document summarization"""
        try:
            documents = []
            
            # Get the main text content
            main_text = processed_content.get('cleaned_text', processed_content.get('raw_text', ''))
            
            if not main_text:
                return ["No content available for summarization"]
            
            # Split into manageable chunks for summarization
            # Use sections if available, otherwise split by length
            structured = processed_content.get('structured_content', {})
            sections = structured.get('sections', [])
            
            if sections and len(sections) > 1:
                # Use sections
                for section in sections:
                    title = section.get('title', '').strip()
                    content = section.get('content', '').strip()
                    
                    if content and len(content) > 100:  # Skip very short sections
                        section_text = f"{title}\n\n{content}" if title else content
                        documents.append(section_text)
            else:
                # Split by length (chunk size for summarization)
                chunk_size = 2000  # Characters per chunk
                chunks = []
                
                paragraphs = main_text.split('\n\n')
                current_chunk = ""
                
                for paragraph in paragraphs:
                    if len(current_chunk) + len(paragraph) < chunk_size:
                        current_chunk += paragraph + "\n\n"
                    else:
                        if current_chunk.strip():
                            chunks.append(current_chunk.strip())
                        current_chunk = paragraph + "\n\n"
                
                # Add remaining chunk
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                
                documents = chunks if chunks else [main_text]
            
            # Ensure we have at least one document
            if not documents:
                documents = [main_text]
            
            # Limit the number of documents to avoid overwhelming the summarizer
            if len(documents) > 10:
                # Keep the first and last few sections, and sample from the middle
                documents = documents[:3] + documents[-3:] + documents[3:-3:2]
            
            return documents[:10]  # Maximum 10 chunks
            
        except Exception as e:
            logger.error(f"Summarization preparation failed: {e}")
            # Fallback: return the raw text split into reasonable chunks
            main_text = processed_content.get('cleaned_text', processed_content.get('raw_text', ''))
            if main_text:
                # Simple splitting by length
                chunk_size = 2000
                chunks = [main_text[i:i+chunk_size] for i in range(0, len(main_text), chunk_size)]
                return chunks[:5]  # Limit to 5 chunks
            else:
                return ["Error: No content available for summarization"]
    
    def get_document_info(self, processed_content: Dict) -> Dict[str, Any]:
        """Get summary information about the processed document"""
        try:
            metrics = processed_content.get('metrics', {})
            structured = processed_content.get('structured_content', {})
            
            info = {
                'filename': processed_content.get('filename', 'Unknown'),
                'file_type': processed_content.get('file_type', 'Unknown'),
                'file_size_bytes': processed_content.get('file_size', 0),
                'processing_success': processed_content.get('success', False),
                'word_count': metrics.get('word_count', 0),
                'character_count': metrics.get('character_count', 0),
                'sentence_count': metrics.get('sentence_count', 0),
                'paragraph_count': metrics.get('paragraph_count', 0),
                'reading_time_minutes': metrics.get('reading_time_minutes', 0),
                'complexity_score': metrics.get('complexity_score', 0),
                'sections_found': len(structured.get('sections', [])),
                'entities_found': {
                    'numbers': len(structured.get('entities', {}).get('numbers', [])),
                    'dates': len(structured.get('entities', {}).get('dates', [])),
                    'urls': len(structured.get('entities', {}).get('urls', [])),
                    'emails': len(structured.get('entities', {}).get('emails', []))
                },
                'key_phrases_count': len(structured.get('key_phrases', [])),
                'processed_at': processed_content.get('processed_at', datetime.now().isoformat())
            }
            
            return info
            
        except Exception as e:
            logger.error(f"Document info extraction failed: {e}")
            return {
                'filename': processed_content.get('filename', 'Unknown'),
                'error': str(e),
                'processing_success': False
            }